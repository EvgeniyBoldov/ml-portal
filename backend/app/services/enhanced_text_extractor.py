from __future__ import annotations

from dataclasses import dataclass
from typing import Any, Dict, List, Tuple, Optional
from io import BytesIO
import csv
import json
import base64

# External deps (add to requirements):
# - pdfminer.six
# - PyPDF2 (optional fallback/meta)
# - python-docx
# - openpyxl
# - charset-normalizer
# - pdfplumber (for tables)
# - pdf2image (for OCR)
# - pytesseract (for OCR)
# - mammoth (for DOCX)
# - pandas (for CSV/Excel)

@dataclass
class TableData:
    name: str
    csv_data: str
    rows: int
    cols: int

@dataclass
class ExtractResult:
    text: str
    kind: str
    meta: Dict[str, Any]
    warnings: List[str]
    tables: List[TableData]

    def to_json(self) -> str:
        return json.dumps(
            {
                "text": self.text,
                "tables": [{"name": t.name, "csv": t.csv_data, "rows": t.rows, "cols": t.cols} for t in self.tables],
                "meta": self.meta,
                "extractor": self.kind,
                "warnings": self.warnings,
            },
            ensure_ascii=False
        )

def extract_text_enhanced(content: bytes, filename: str) -> ExtractResult:
    """
    Enhanced text extraction with OCR fallback and table support
    """
    if not content:
        return ExtractResult("", "empty", {}, [], [])
    
    ext = filename.lower().split('.')[-1] if '.' in filename else ''
    warnings = []
    tables = []
    
    # Text files
    if ext in ('txt', 'md', 'rtf'):
        return _extract_text_file(content, ext, warnings)
    
    # CSV files
    elif ext == 'csv':
        return _extract_csv_file(content, warnings)
    
    # PDF files
    elif ext == 'pdf':
        return _extract_pdf_enhanced(content, warnings)
    
    # DOCX files
    elif ext == 'docx':
        return _extract_docx_enhanced(content, warnings)
    
    # XLSX files
    elif ext == 'xlsx':
        return _extract_xlsx_enhanced(content, warnings)
    
    # DOC files (legacy)
    elif ext == 'doc':
        return _extract_doc_file(content, warnings)
    
    # Unknown format - try as text
    else:
        warnings.append(f"Unknown extension '.{ext}', treated as text.")
        return _extract_text_file(content, 'txt', warnings)

def _extract_text_file(content: bytes, ext: str, warnings: List[str]) -> ExtractResult:
    """Extract text from plain text files"""
    try:
        # Try to detect encoding
        import charset_normalizer
        detected = charset_normalizer.detect(content)
        encoding = detected.get('encoding', 'utf-8')
        confidence = detected.get('confidence', 0.0)
        
        if confidence < 0.7:
            warnings.append(f"Low confidence encoding detection: {encoding} ({confidence:.2f})")
        
        text = content.decode(encoding, errors='replace')
        
        # Handle RTF
        if ext == 'rtf':
            text = _clean_rtf(text)
        
        return ExtractResult(
            text=text,
            kind=f"txt({encoding})",
            meta={"encoding": encoding, "confidence": confidence},
            warnings=warnings,
            tables=[]
        )
    except Exception as e:
        warnings.append(f"Text extraction failed: {e}")
        return ExtractResult(
            text=content.decode('utf-8', errors='replace'),
            kind="txt(utf_8_fallback)",
            meta={"encoding": "utf-8", "error": str(e)},
            warnings=warnings,
            tables=[]
        )

def _extract_csv_file(content: bytes, warnings: List[str]) -> ExtractResult:
    """Extract text and tables from CSV files"""
    try:
        import pandas as pd
        import io
        
        # Detect encoding
        import charset_normalizer
        detected = charset_normalizer.detect(content)
        encoding = detected.get('encoding', 'utf-8')
        
        # Read CSV
        df = pd.read_csv(io.BytesIO(content), encoding=encoding)
        
        # Convert to text
        text = df.to_string(index=False)
        
        # Create table data
        csv_data = df.to_csv(index=False)
        table = TableData(
            name="main_table",
            csv_data=csv_data,
            rows=len(df),
            cols=len(df.columns)
        )
        
        return ExtractResult(
            text=text,
            kind="csv(enhanced)",
            meta={"encoding": encoding, "rows": len(df), "cols": len(df.columns)},
            warnings=warnings,
            tables=[table]
        )
    except Exception as e:
        warnings.append(f"Enhanced CSV extraction failed: {e}")
        # Fallback to simple CSV
        return _extract_simple_csv(content, warnings)

def _extract_simple_csv(content: bytes, warnings: List[str]) -> ExtractResult:
    """Simple CSV extraction fallback"""
    try:
        import csv
        import io
        
        # Detect encoding
        import charset_normalizer
        detected = charset_normalizer.detect(content)
        encoding = detected.get('encoding', 'utf-8')
        
        # Read CSV
        text_content = content.decode(encoding, errors='replace')
        csv_reader = csv.reader(io.StringIO(text_content))
        rows = list(csv_reader)
        
        # Format as text
        text = '\n'.join(['\t'.join(row) for row in rows])
        
        # Create table data
        csv_data = text_content
        table = TableData(
            name="main_table",
            csv_data=csv_data,
            rows=len(rows),
            cols=len(rows[0]) if rows else 0
        )
        
        return ExtractResult(
            text=text,
            kind="csv(simple)",
            meta={"encoding": encoding, "rows": len(rows)},
            warnings=warnings,
            tables=[table]
        )
    except Exception as e:
        warnings.append(f"CSV extraction failed: {e}")
        return ExtractResult(
            text=content.decode('utf-8', errors='replace'),
            kind="csv(fallback)",
            meta={"encoding": "utf-8", "error": str(e)},
            warnings=warnings,
            tables=[]
        )

def _extract_pdf_enhanced(content: bytes, warnings: List[str]) -> ExtractResult:
    """Enhanced PDF extraction with OCR fallback and table support"""
    text = ""
    tables = []
    meta = {}
    
    # Try pdfplumber first (better for tables)
    try:
        import pdfplumber
        with pdfplumber.open(BytesIO(content)) as pdf:
            pages_text = []
            all_tables = []
            
            for page_num, page in enumerate(pdf.pages):
                # Extract text
                page_text = page.extract_text()
                if page_text:
                    pages_text.append(page_text)
                
                # Extract tables
                page_tables = page.extract_tables()
                for table_num, table in enumerate(page_tables):
                    if table and len(table) > 1:  # Skip empty tables
                        # Convert to CSV
                        csv_data = '\n'.join([','.join([str(cell or '') for cell in row]) for row in table])
                        table_data = TableData(
                            name=f"page_{page_num+1}_table_{table_num+1}",
                            csv_data=csv_data,
                            rows=len(table),
                            cols=len(table[0]) if table else 0
                        )
                        all_tables.append(table_data)
            
            text = '\n\n'.join(pages_text)
            tables = all_tables
            meta = {"pages": len(pdf.pages), "tables_found": len(tables)}
            
            if text.strip():
                return ExtractResult(
                    text=text,
                    kind="pdf(pdfplumber)",
                    meta=meta,
                    warnings=warnings,
                    tables=tables
                )
    except Exception as e:
        warnings.append(f"PDF extraction via pdfplumber failed: {e}")
    
    # Try pdfminer.six
    try:
        from pdfminer.high_level import extract_text
        text = extract_text(BytesIO(content))
        if text.strip():
            return ExtractResult(
                text=text,
                kind="pdf(pdfminer)",
                meta={"pages": 1, "method": "pdfminer"},
                warnings=warnings,
                tables=tables
            )
    except Exception as e:
        warnings.append(f"PDF extraction via pdfminer failed: {e}")
    
    # Try PyPDF2
    try:
        import PyPDF2
        pdf_reader = PyPDF2.PdfReader(BytesIO(content))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        
        if text.strip():
            return ExtractResult(
                text=text,
                kind="pdf(pypdf2)",
                meta={"pages": len(pdf_reader.pages), "method": "pypdf2"},
                warnings=warnings,
                tables=tables
            )
    except Exception as e:
        warnings.append(f"PDF extraction via PyPDF2 failed: {e}")
    
    # OCR fallback
    try:
        return _extract_pdf_ocr(content, warnings)
    except Exception as e:
        warnings.append(f"OCR extraction failed: {e}")
    
    # Final fallback
    return ExtractResult(
        text="",
        kind="pdf(failed)",
        meta={"pages": 0, "error": "All extraction methods failed"},
        warnings=warnings,
        tables=[]
    )

def _extract_pdf_ocr(content: bytes, warnings: List[str]) -> ExtractResult:
    """OCR extraction for scanned PDFs"""
    try:
        from pdf2image import convert_from_bytes
        import pytesseract
        
        # Convert PDF to images
        images = convert_from_bytes(content, dpi=300)
        
        text_parts = []
        for i, image in enumerate(images):
            # OCR each page
            page_text = pytesseract.image_to_string(image, lang='rus+eng')
            text_parts.append(page_text)
        
        text = '\n\n'.join(text_parts)
        
        return ExtractResult(
            text=text,
            kind="pdf(ocr)",
            meta={"pages": len(images), "method": "ocr", "lang": "rus+eng"},
            warnings=warnings,
            tables=[]
        )
    except Exception as e:
        raise Exception(f"OCR processing failed: {e}")

def _extract_docx_enhanced(content: bytes, warnings: List[str]) -> ExtractResult:
    """Enhanced DOCX extraction with table support"""
    try:
        from docx import Document
        from docx.table import Table
        
        doc = Document(BytesIO(content))
        
        # Extract text
        text_parts = []
        tables = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extract tables
        for table_num, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            
            if table_data:
                # Convert to CSV
                csv_data = '\n'.join([','.join(row) for row in table_data])
                table_obj = TableData(
                    name=f"table_{table_num+1}",
                    csv_data=csv_data,
                    rows=len(table_data),
                    cols=len(table_data[0]) if table_data else 0
                )
                tables.append(table_obj)
        
        text = '\n'.join(text_parts)
        
        return ExtractResult(
            text=text,
            kind="docx(enhanced)",
            meta={"tables_found": len(tables)},
            warnings=warnings,
            tables=tables
        )
    except Exception as e:
        warnings.append(f"Enhanced DOCX extraction failed: {e}")
        # Fallback to mammoth
        return _extract_docx_mammoth(content, warnings)

def _extract_docx_mammoth(content: bytes, warnings: List[str]) -> ExtractResult:
    """DOCX extraction using mammoth"""
    try:
        import mammoth
        
        result = mammoth.extract_raw_text(BytesIO(content))
        text = result.value
        
        return ExtractResult(
            text=text,
            kind="docx(mammoth)",
            meta={"method": "mammoth"},
            warnings=warnings,
            tables=[]
        )
    except Exception as e:
        warnings.append(f"DOCX extraction via mammoth failed: {e}")
        return ExtractResult(
            text="",
            kind="docx(failed)",
            meta={"error": str(e)},
            warnings=warnings,
            tables=[]
        )

def _extract_xlsx_enhanced(content: bytes, warnings: List[str]) -> ExtractResult:
    """Enhanced XLSX extraction with table support"""
    try:
        import openpyxl
        import pandas as pd
        import io
        
        # Load workbook
        workbook = openpyxl.load_workbook(BytesIO(content))
        
        text_parts = []
        tables = []
        
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            
            # Extract text from sheet
            sheet_text = []
            for row in sheet.iter_rows(values_only=True):
                row_text = [str(cell or '') for cell in row]
                if any(cell.strip() for cell in row_text):
                    sheet_text.append('\t'.join(row_text))
            
            if sheet_text:
                text_parts.append(f"Sheet: {sheet_name}\n" + '\n'.join(sheet_text))
            
            # Create table data
            if sheet_text:
                # Convert to DataFrame for CSV
                df = pd.read_excel(BytesIO(content), sheet_name=sheet_name)
                csv_data = df.to_csv(index=False)
                
                table = TableData(
                    name=sheet_name,
                    csv_data=csv_data,
                    rows=len(df),
                    cols=len(df.columns)
                )
                tables.append(table)
        
        text = '\n\n'.join(text_parts)
        
        return ExtractResult(
            text=text,
            kind="xlsx(enhanced)",
            meta={"sheets": len(workbook.sheetnames), "tables_found": len(tables)},
            warnings=warnings,
            tables=tables
        )
    except Exception as e:
        warnings.append(f"Enhanced XLSX extraction failed: {e}")
        return ExtractResult(
            text="",
            kind="xlsx(failed)",
            meta={"error": str(e)},
            warnings=warnings,
            tables=[]
        )

def _extract_doc_file(content: bytes, warnings: List[str]) -> ExtractResult:
    """Extract from legacy DOC files"""
    try:
        import mammoth
        
        result = mammoth.extract_raw_text(BytesIO(content))
        text = result.value
        
        return ExtractResult(
            text=text,
            kind="doc(mammoth)",
            meta={"method": "mammoth"},
            warnings=warnings,
            tables=[]
        )
    except Exception as e:
        warnings.append(f"DOC extraction failed: {e}")
        return ExtractResult(
            text="",
            kind="doc(failed)",
            meta={"error": str(e)},
            warnings=warnings,
            tables=[]
        )

def _clean_rtf(text: str) -> str:
    """Basic RTF cleaning"""
    import re
    
    # Remove RTF control codes
    text = re.sub(r'\\[a-z]+\d*\s?', '', text)
    text = re.sub(r'[{}]', '', text)
    text = re.sub(r'\s+', ' ', text)
    
    return text.strip()
