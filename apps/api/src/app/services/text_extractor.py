from __future__ import annotations

from dataclasses import dataclass
from typing import Any, Dict, List, Tuple
from io import BytesIO
import csv
import json

# External deps (add to requirements):
# - pdfminer.six
# - PyPDF2 (optional fallback/meta)
# - python-docx
# - openpyxl
# - charset-normalizer
#
# We import lazily inside functions to avoid import errors when a format isn't used.

@dataclass
class ExtractResult:
    text: str
    kind: str
    meta: Dict[str, Any]
    warnings: List[str]

    def to_json(self) -> str:
        return json.dumps(
            {
                "text": self.text,
                "type": "text",
                "extractor": self.kind,
                "meta": self.meta,
                "warnings": self.warnings,
            },
            ensure_ascii=False,
        )


def _detect_ext(filename: str) -> str:
    name = (filename or "").lower()
    if "." not in name:
        return ""
    return name[name.rfind(".") + 1 :]  # ext without dot


def _decode_best_effort(data: bytes) -> Tuple[str, str, List[str]]:
    """Decode bytes to str using charset-normalizer (fallback to utf-8)."""
    warnings: List[str] = []
    try:
        from charset_normalizer import from_bytes as cn_from_bytes  # type: ignore
        res = cn_from_bytes(data).best()
        if res is not None:
            return str(res), (res.encoding or "utf-8"), warnings
        warnings.append("charset-normalizer: no best match, fallback to utf-8.")
        return data.decode("utf-8", errors="replace"), "utf-8", warnings
    except Exception as e:  # optional dep may be absent in tests
        warnings.append(f"charset-normalizer unavailable ({e!r}); fallback to utf-8.")
        return data.decode("utf-8", errors="replace"), "utf-8", warnings


def _extract_pdf(data: bytes) -> ExtractResult:
    warnings: List[str] = []
    text = ""
    pages = 0
    try:
        # pdfminer.six is more accurate for text PDFs than PyPDF2
        from pdfminer.high_level import extract_text as pdf_extract_text  # type: ignore
        text = pdf_extract_text(BytesIO(data)) or ""
        # Try PyPDF2 for meta/page count
        try:
            import PyPDF2  # type: ignore
            r = PyPDF2.PdfReader(BytesIO(data))
            pages = len(r.pages)
        except Exception:
            pages = 0
        if not text.strip():
            warnings.append("PDF appears to have no extractable text (maybe scanned). Consider OCR later.")
    except Exception as e:
        warnings.append(f"PDF extraction failed via pdfminer: {e!r}")
        # last resort: try PyPDF2.extract_text
        try:
            import PyPDF2  # type: ignore
            r = PyPDF2.PdfReader(BytesIO(data))
            pages = len(r.pages)
            text = ""
            for p in r.pages:
                try:
                    text += (p.extract_text() or "") + "\n"
                except Exception:
                    continue
            if not text.strip():
                warnings.append("PyPDF2 yielded empty text. Consider OCR.")
        except Exception as e2:
            warnings.append(f"PyPDF2 fallback failed: {e2!r}")
            text = ""
    return ExtractResult(text=text, kind="pdf", meta={"pages": pages}, warnings=warnings)


def _extract_docx(data: bytes) -> ExtractResult:
    warnings: List[str] = []
    text = ""
    try:
        from docx import Document  # type: ignore
        doc = Document(BytesIO(data))
        parts: List[str] = []
        # paragraphs
        for p in doc.paragraphs:
            if p.text:
                parts.append(p.text)
        # tables
        for t in doc.tables:
            for row in t.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    parts.append("\t".join(cells))
        text = "\n".join(parts).strip()
        if not text:
            warnings.append("DOCX parsed but no visible text was found.")
    except Exception as e:
        warnings.append(f"DOCX extraction failed: {e!r}")
        text = ""
    return ExtractResult(text=text, kind="docx", meta={}, warnings=warnings)


def _extract_txt(data: bytes) -> ExtractResult:
    text, enc, warn = _decode_best_effort(data)
    return ExtractResult(text=text, kind=f"txt({enc})", meta={"encoding": enc}, warnings=warn)


def _extract_csv(data: bytes) -> ExtractResult:
    text, enc, warn = _decode_best_effort(data)
    # Try sniffing CSV dialect; render as TSV-like plain text
    out_lines: List[str] = []
    try:
        sniffer = csv.Sniffer()
        dialect = sniffer.sniff(text[:10000])
    except Exception:
        dialect = csv.excel
    reader = csv.reader(text.splitlines(), dialect=dialect)
    for row in reader:
        out_lines.append("\t".join("" if v is None else str(v) for v in row))
    return ExtractResult(text="\n".join(out_lines), kind=f"csv({enc})", meta={"encoding": enc}, warnings=warn)


def _extract_xlsx(data: bytes) -> ExtractResult:
    warnings: List[str] = []
    text = ""
    try:
        from openpyxl import load_workbook  # type: ignore
        wb = load_workbook(BytesIO(data), data_only=True, read_only=True)
        out_lines: List[str] = []
        for ws in wb.worksheets:
            out_lines.append(f"# Sheet: {ws.title}")
            for row in ws.iter_rows(values_only=True):
                out_lines.append("\t".join("" if v is None else str(v) for v in row))
            out_lines.append("")
        text = "\n".join(out_lines).strip()
    except Exception as e:
        warnings.append(f"XLSX extraction failed: {e!r}")
        text = ""
    return ExtractResult(text=text, kind="xlsx", meta={}, warnings=warnings)


def extract_text(data: bytes, filename: str) -> ExtractResult:
    """
    Universal text extractor for: PDF, DOCX, TXT, CSV, XLSX.
    Returns ExtractResult(text, kind, meta, warnings).
    """
    ext = _detect_ext(filename)
    if ext in ("pdf",):
        return _extract_pdf(data)
    if ext in ("docx",):
        return _extract_docx(data)
    if ext in ("txt", "log", ""):
        return _extract_txt(data)
    if ext in ("csv",):
        return _extract_csv(data)
    if ext in ("xlsx",):
        return _extract_xlsx(data)
    # Unknown → try text decode
    res = _extract_txt(data)
    res.warnings.append(f"Unknown extension '.{ext}', treated as text.")
    return res
