"""
TemplateAnalyzeService — parse uploaded templates (Excel, Word, text) to extract
structure: placeholders, tables, headers, title, version, and a draft fill schema.
"""
from __future__ import annotations

import io
import re
from typing import Any, Dict, List, Optional

from app.core.logging import get_logger

logger = get_logger(__name__)

_PLACEHOLDER_RE = re.compile(r"\{\{([^}]+)\}\}")


def _extract_title_and_version_from_text(texts: List[str]) -> tuple[Optional[str], Optional[str]]:
    title = None
    version = None
    for raw in texts:
        line = raw.strip()
        if not line:
            continue
        lower = line.lower()
        if title is None and len(line) > 3:
            title = line
        if version is None:
            for pattern in (r"версия[:\s]*([\d.]+)", r"version[:\s]*([\d.]+)", r"v\.?\s*([\d.]+)"):
                m = re.search(pattern, lower)
                if m:
                    version = m.group(1)
                    break
        if title and version:
            break
    return title, version


class TemplateAnalyzeService:
    """Analyze a template file and return structured metadata + draft schema."""

    async def analyze_bytes(self, content: bytes, filename: str) -> dict:
        schema_payload = await self.generate_schema(content, filename)
        description_payload = await self.generate_description(content, filename)
        return {
            **schema_payload,
            **description_payload,
        }

    async def generate_schema(self, content: bytes, filename: str) -> dict:
        ext = (filename.split(".")[-1] or "").lower()
        if ext in ("xlsx", "xlsm"):
            analysis = self._analyze_excel(content, filename)
        elif ext == "docx":
            analysis = self._analyze_docx(content, filename)
        else:
            analysis = self._analyze_text(content, filename)
        return {
            "title": analysis.get("title"),
            "version": analysis.get("version"),
            "draft_schema": analysis.get("draft_schema"),
        }

    async def generate_description(self, content: bytes, filename: str) -> dict:
        ext = (filename.split(".")[-1] or "").lower()
        if ext in ("xlsx", "xlsm"):
            analysis = self._analyze_excel(content, filename)
        elif ext == "docx":
            analysis = self._analyze_docx(content, filename)
        else:
            analysis = self._analyze_text(content, filename)
        return {
            "title": analysis.get("title"),
            "version": analysis.get("version"),
            "description": analysis.get("description"),
        }

    def _analyze_excel(self, content: bytes, filename: str) -> dict:
        try:
            import openpyxl
        except ImportError as exc:
            raise RuntimeError("openpyxl is required for Excel template analysis") from exc

        wb = openpyxl.load_workbook(io.BytesIO(content), data_only=True)
        sheets_meta = []
        all_placeholders = []
        first_texts = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            cells = []
            sheet_texts = []
            for row in ws.iter_rows():
                for cell in row:
                    if cell.value is not None:
                        val = str(cell.value).strip()
                        if val:
                            cells.append({
                                "row": cell.row,
                                "col": cell.column,
                                "value": val,
                                "coordinate": cell.coordinate,
                            })
                            sheet_texts.append(val)
                            for ph in _PLACEHOLDER_RE.findall(val):
                                all_placeholders.append({
                                    "placeholder": ph.strip(),
                                    "sheet": sheet_name,
                                    "coordinate": cell.coordinate,
                                })
            first_texts.extend(sheet_texts[:5])

            # Detect tables: find rectangular regions with dense data
            # Simple heuristic: if >3 rows have same number of non-empty columns → table
            row_vals = {}
            for c in cells:
                row_vals.setdefault(c["row"], []).append(c["value"])
            # Find header row: first row with >1 non-empty cells
            header_row = None
            sorted_rows = sorted(row_vals.keys())
            for r in sorted_rows:
                if len(row_vals[r]) > 1:
                    header_row = r
                    break

            sheets_meta.append({
                "name": sheet_name,
                "cell_count": len(cells),
                "header_row": header_row,
                "placeholders": len([p for p in all_placeholders if p["sheet"] == sheet_name]),
            })

        title, version = _extract_title_and_version_from_text(first_texts)

        draft_schema = {
            "format": "excel",
            "sheets": sheets_meta,
            "placeholders": all_placeholders,
        }

        return {
            "title": title,
            "version": version,
            "draft_schema": draft_schema,
            "kind_hint": self._infer_kind(title, all_placeholders),
            "description": self._build_description(filename, "excel", title, all_placeholders),
        }

    def _analyze_docx(self, content: bytes, filename: str) -> dict:
        try:
            import docx
        except ImportError as exc:
            raise RuntimeError("python-docx is required for Word template analysis") from exc

        doc = docx.Document(io.BytesIO(content))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
        all_placeholders = []
        for idx, para in enumerate(paragraphs):
            for ph in _PLACEHOLDER_RE.findall(para):
                all_placeholders.append({
                    "placeholder": ph.strip(),
                    "paragraph_index": idx,
                    "context": para[:200],
                })

        tables_meta = []
        for t_idx, table in enumerate(doc.tables):
            rows = len(table.rows)
            cols = len(table.columns)
            tables_meta.append({"index": t_idx, "rows": rows, "columns": cols})

        title, version = _extract_title_and_version_from_text(paragraphs[:10])

        draft_schema = {
            "format": "docx",
            "paragraph_count": len(paragraphs),
            "tables": tables_meta,
            "placeholders": all_placeholders,
        }

        return {
            "title": title,
            "version": version,
            "draft_schema": draft_schema,
            "kind_hint": self._infer_kind(title, all_placeholders),
            "description": self._build_description(filename, "docx", title, all_placeholders),
        }

    def _analyze_text(self, content: bytes, filename: str) -> dict:
        try:
            text = content.decode("utf-8")
        except UnicodeDecodeError:
            text = content.decode("latin-1")
        lines = [l.strip() for l in text.splitlines() if l.strip()]
        all_placeholders = []
        for idx, line in enumerate(lines):
            for ph in _PLACEHOLDER_RE.findall(line):
                all_placeholders.append({
                    "placeholder": ph.strip(),
                    "line_index": idx,
                    "context": line[:200],
                })
        title, version = _extract_title_and_version_from_text(lines[:10])
        draft_schema = {
            "format": "text",
            "line_count": len(lines),
            "placeholders": all_placeholders,
        }
        return {
            "title": title,
            "version": version,
            "draft_schema": draft_schema,
            "kind_hint": self._infer_kind(title, all_placeholders),
            "description": self._build_description(filename, "text", title, all_placeholders),
        }

    @staticmethod
    def _infer_kind(title: Optional[str], placeholders: List[dict]) -> Optional[str]:
        if not title:
            return None
        t = title.lower()
        if any(w in t for w in ("заявк", "request", "запрос")):
            return "request"
        if any(w in t for w in ("план", "plan")):
            return "plan"
        if any(w in t for w in ("конфиг", "config", "настройк")):
            return "config"
        if any(w in t for w in ("акт", "act", "report")):
            return "report"
        if placeholders:
            return "form"
        return None

    @staticmethod
    def _build_description(
        filename: str,
        fmt: str,
        title: Optional[str],
        placeholders: List[dict],
    ) -> str:
        normalized_placeholders: list[str] = []
        for item in placeholders:
            placeholder = str(item.get("placeholder") or "").strip()
            if placeholder and placeholder not in normalized_placeholders:
                normalized_placeholders.append(placeholder)

        subject = title or filename
        if normalized_placeholders:
            placeholder_list = ", ".join(normalized_placeholders[:12])
            return (
                f"Template '{subject}' ({fmt}) with placeholders: {placeholder_list}."
            )
        return f"Template '{subject}' ({fmt}) without detected placeholders."
