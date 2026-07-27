"""
TemplateLayoutParser — deterministic raw layout extraction (S1).

Parses template files (Excel / Word / plain-text) and returns a ``RawLayout``
dataclass that captures:

- All ``{{token}}`` occurrences with their positions.
- Document-level structure (sheets / paragraphs / lines).
- Table regions (Excel openpyxl tables, docx tables, heuristic dense regions).
- Marker-loop candidates: rows/columns that contain at least one dotted token
  ``{{table.col}}``.
- Block fences: ``{{#key}} … {{/key}}`` pairs in docx/text.
- Title and version hints extracted from leading non-placeholder text.

This module has **no LLM calls, no DB access, no async I/O** — it is a pure
bytes → data-structure transformation and must remain deterministic.
"""
from __future__ import annotations

import re
from dataclasses import asdict, dataclass, field
from typing import Any, Dict, List, Optional, Tuple

# ---------------------------------------------------------------------------
# Regex patterns
# ---------------------------------------------------------------------------

# Matches any placeholder expression: {{name}}, {{author.tel:int(10)}}
_TOKEN_RE = re.compile(r"\{\{([^{}]+)\}\}")
# Open/close fences: {{#key}} / {{/key}}
_FENCE_OPEN_RE = re.compile(r"\{\{#([A-Za-z0-9_.\-]+)\}\}")
_FENCE_CLOSE_RE = re.compile(r"\{\{/([A-Za-z0-9_.\-]+)\}\}")
# Version hint
_VERSION_RE = re.compile(
    r"(?:версия|version|v\.?)[:\s]*([\d]+(?:\.[\d]+)*)", re.IGNORECASE
)


# ---------------------------------------------------------------------------
# Data structures
# ---------------------------------------------------------------------------


@dataclass
class TokenOccurrence:
    """A single ``{{token}}`` occurrence in the source."""
    token: str
    table_prefix: Optional[str]
    column_key: Optional[str]
    location: Dict[str, Any]
    placeholder: Optional[str] = None
    hint_type: Optional[str] = None
    hint_args: Optional[str] = None
    segments: List[Dict[str, Any]] = field(default_factory=list)
    repeat_root: Optional[str] = None
    params: Dict[str, Any] = field(default_factory=dict)


@dataclass
class PathSegment:
    """Parsed path segment from the placeholder contract."""
    key: str
    params: Dict[str, Any] = field(default_factory=dict)
    repeat: bool = False


@dataclass
class SchemaNode:
    """Merged schema node assembled from placeholder segments."""
    key: str
    path: str
    params: Dict[str, Any] = field(default_factory=dict)
    repeat: bool = False
    placeholder: Optional[str] = None
    location: Dict[str, Any] = field(default_factory=dict)
    hint_type: Optional[str] = None
    hint_args: Optional[str] = None
    anchors: List[Dict[str, Any]] = field(default_factory=list)
    children: List["SchemaNode"] = field(default_factory=list)


@dataclass
class TableRegion:
    """A detected repeatable region in the document."""
    region_id: str
    location: Dict[str, Any]           # {sheet, row_start, row_end, col_start, col_end} etc.
    # Marker-loop info (if any column token found in the region)
    loop_tokens: List[str] = field(default_factory=list)   # e.g. ["{{items.name}}", "{{items.qty}}"]
    marker_columns: Dict[str, int] = field(default_factory=dict)
    loop_prefix: Optional[str] = None                       # common table prefix, e.g. "items"
    # Structural anchor hint
    header_row: Optional[List[str]] = field(default_factory=list)  # text of header cells
    template_row_index: Optional[int] = None                        # 0-based row index after header
    orientation: str = "vertical"   # "vertical" | "horizontal"


@dataclass
class FenceBlock:
    """A ``{{#key}} … {{/key}}`` block (docx/text)."""
    key: str
    open_position: Dict[str, Any]
    close_position: Optional[Dict[str, Any]] = None


@dataclass
class RawLayout:
    """Complete raw layout extracted from a template file."""
    format: str                                    # "excel" | "docx" | "text"
    title: Optional[str]
    version: Optional[str]
    tokens: List[TokenOccurrence] = field(default_factory=list)
    table_regions: List[TableRegion] = field(default_factory=list)
    fence_blocks: List[FenceBlock] = field(default_factory=list)
    sheets: List[str] = field(default_factory=list)   # Excel sheet names
    # Aggregated unique token keys (scalar and table-prefixed)
    scalar_keys: List[str] = field(default_factory=list)   # non-dotted keys
    table_prefixes: List[str] = field(default_factory=list)  # dotted key table-parts
    schema_roots: List[SchemaNode] = field(default_factory=list)
    # Raw text lines / paragraph texts — useful for LLM schema builder
    text_lines: List[str] = field(default_factory=list)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


class TemplateLayoutParser:
    """Parse a template file and return a ``RawLayout`` (no LLM, no I/O)."""

    def parse(self, content: bytes, filename: str) -> RawLayout:
        ext = _ext(filename)
        if ext not in ("xlsx", "xlsm"):
            raise ValueError("Only .xlsx and .xlsm templates are supported")
        return self._parse_excel(content, filename)

    # ------------------------------------------------------------------
    # Excel
    # ------------------------------------------------------------------

    def _parse_excel(self, content: bytes, filename: str) -> RawLayout:
        import io

        try:
            import openpyxl
        except ImportError as exc:
            raise RuntimeError("openpyxl is required for Excel template parsing") from exc

        # For template analysis we need the source cell text, including formulas
        # that may produce placeholder strings. ``data_only=True`` drops formula
        # bodies and often returns ``None`` when cached values are absent.
        wb = openpyxl.load_workbook(
            io.BytesIO(content), data_only=False, keep_vba=filename.lower().endswith(".xlsm")
        )
        tokens: List[TokenOccurrence] = []
        table_regions: List[TableRegion] = []
        text_lines: List[str] = []
        first_texts: List[str] = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            # Collect all cell data: {row: {col: value}}
            row_map: Dict[int, Dict[int, str]] = {}
            for row in ws.iter_rows():
                for cell in row:
                    if cell.value is None:
                        continue
                    val = str(cell.value).strip()
                    if not val:
                        continue
                    row_map.setdefault(cell.row, {})[cell.column] = val
                    text_lines.append(val)

                    for m in _TOKEN_RE.finditer(val):
                        parsed = _parse_placeholder_expr(m.group(1))
                        if not parsed:
                            continue
                        key, hint_type, hint_args, spec = parsed
                        table_prefix, col_key = _resolve_repeat_path(spec)
                        tokens.append(TokenOccurrence(
                            token=key,
                            table_prefix=table_prefix,
                            column_key=col_key,
                            location={
                                "sheet": sheet_name,
                                "row": cell.row,
                                "col": cell.column,
                                "coordinate": cell.coordinate,
                                "source_text": val,
                            },
                            placeholder=m.group(0),
                            hint_type=hint_type,
                            hint_args=hint_args,
                            segments=spec["segments"],
                            repeat_root=spec["repeat_root"],
                            params=spec["params"],
                        ))

            if first_texts == [] and text_lines:
                first_texts = text_lines[:10]

            # Detect table regions for this sheet
            table_regions.extend(
                self._detect_excel_table_regions(ws, sheet_name, row_map)
            )

        sheet_names = list(wb.sheetnames)
        wb.close()

        title, version = _extract_title_version(first_texts)
        scalar_keys, table_prefixes = _aggregate_keys(tokens, table_regions)

        return RawLayout(
            format="excel",
            title=title,
            version=version,
            tokens=tokens,
            table_regions=table_regions,
            sheets=sheet_names,
            scalar_keys=scalar_keys,
            table_prefixes=table_prefixes,
            schema_roots=_build_schema_roots(tokens, table_regions),
            text_lines=text_lines[:500],  # cap for LLM builder
        )

    def _detect_excel_table_regions(
        self,
        ws: Any,
        sheet_name: str,
        row_map: Dict[int, Dict[int, str]],
    ) -> List[TableRegion]:
        """Detect repeat regions solely from explicit technical marker rows."""
        regions: List[TableRegion] = []
        sorted_rows = sorted(row_map.keys())

        # --- Marker detection ---
        prefix_to_rows: Dict[str, List[int]] = {}
        for r in sorted_rows:
            row_vals = row_map[r]
            prefixes_in_row: Dict[str, List[str]] = {}
            for val in row_vals.values():
                for m in _TOKEN_RE.finditer(val):
                    parsed = _parse_placeholder_expr(m.group(1))
                    if not parsed:
                        continue
                    _, _, _, spec = parsed
                    tp, _ = _resolve_repeat_path(spec)
                    if tp:
                        prefixes_in_row.setdefault(tp, []).append(m.group(0))
            for tp, toks in prefixes_in_row.items():
                prefix_to_rows.setdefault(tp, []).append(r)

        ordered_prefixes = sorted(prefix_to_rows.items(), key=lambda item: item[1][0] if item[1] else 10**9)
        for prefix, marker_rows in ordered_prefixes:
            marker_row = marker_rows[0]
            # Collect all loop_tokens from that row
            loop_tokens: List[str] = []
            marker_columns: Dict[str, int] = {}
            row_cols = sorted(row_map.get(marker_row, {}).keys())
            col_start = row_cols[0] if row_cols else 1
            col_end = row_cols[-1] if row_cols else 1
            for column, val in row_map.get(marker_row, {}).items():
                for m in _TOKEN_RE.finditer(val):
                    parsed = _parse_placeholder_expr(m.group(1))
                    if not parsed:
                        continue
                    _, _, _, spec = parsed
                    tp, _ = _resolve_repeat_path(spec)
                    if tp == prefix:
                        tok = m.group(0)
                        if tok not in loop_tokens:
                            loop_tokens.append(tok)
                        marker_columns[tok] = column
            # Check for optional row above as header
            header_row_idx = marker_row - 1
            header_texts: List[str] = []
            if header_row_idx in row_map:
                header_vals = row_map[header_row_idx]
                header_texts = [header_vals.get(c, "") for c in row_cols]

            if not loop_tokens:
                continue

            regions.append(TableRegion(
                region_id=f"{sheet_name}:marker:{prefix}",
                location={
                    "sheet": sheet_name,
                    "marker_row": marker_row,
                    "col_start": col_start,
                    "col_end": col_end,
                },
                loop_tokens=loop_tokens,
                marker_columns=marker_columns,
                loop_prefix=prefix,
                header_row=header_texts,
                template_row_index=None,
                orientation="vertical",
            ))

        return regions

    def _detect_dense_regions(
        self,
        sheet_name: str,
        row_map: Dict[int, Dict[int, str]],
        sorted_rows: List[int],
    ) -> List[TableRegion]:
        """Find dense rectangular regions (≥2 cols, ≥2 rows) as structural candidates."""
        regions: List[TableRegion] = []
        min_cols = 2
        run_rows: List[int] = []

        def _flush(run: List[int]) -> None:
            if len(run) < 2:
                return
            first_r = run[0]
            last_r = run[-1]
            all_cols: List[int] = []
            for r in run:
                all_cols.extend(row_map[r].keys())
            col_start = min(all_cols) if all_cols else 1
            col_end = max(all_cols) if all_cols else 1
            # Treat first row as potential header (no tokens)
            first_row_vals = list(row_map[first_r].values())
            has_tokens_in_first = any(_TOKEN_RE.search(v) for v in first_row_vals)
            header_texts = first_row_vals if not has_tokens_in_first else []
            template_row_idx = 1 if not has_tokens_in_first else 0

            regions.append(TableRegion(
                region_id=f"{sheet_name}:structural:r{first_r}-{last_r}",
                location={
                    "sheet": sheet_name,
                    "row_start": first_r,
                    "row_end": last_r,
                    "col_start": col_start,
                    "col_end": col_end,
                },
                loop_tokens=[],
                loop_prefix=None,
                header_row=header_texts,
                template_row_index=template_row_idx,
                orientation="vertical",
            ))

        for r in sorted_rows:
            if len(row_map[r]) >= min_cols:
                run_rows.append(r)
            else:
                _flush(run_rows)
                run_rows = []
        _flush(run_rows)
        return regions

    # ------------------------------------------------------------------
    # Docx
    # ------------------------------------------------------------------

    def _parse_docx(self, content: bytes, filename: str) -> RawLayout:
        import io

        try:
            import docx
        except ImportError as exc:
            raise RuntimeError("python-docx is required for Word template parsing") from exc

        doc = docx.Document(io.BytesIO(content))
        tokens: List[TokenOccurrence] = []
        fence_blocks: List[FenceBlock] = []
        text_lines: List[str] = []
        table_regions: List[TableRegion] = []

        open_fences: Dict[str, Dict[str, Any]] = {}

        # Paragraphs
        for p_idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            text_lines.append(text)

            for m in _FENCE_OPEN_RE.finditer(text):
                key = m.group(1)
                open_fences[key] = {"paragraph_index": p_idx, "context": text[:100]}

            for m in _FENCE_CLOSE_RE.finditer(text):
                key = m.group(1)
                open_pos = open_fences.pop(key, None)
                fence_blocks.append(FenceBlock(
                    key=key,
                    open_position=open_pos or {},
                    close_position={"paragraph_index": p_idx, "context": text[:100]},
                ))

            for m in _TOKEN_RE.finditer(text):
                parsed = _parse_placeholder_expr(m.group(1))
                if not parsed:
                    continue
                key, hint_type, hint_args, spec = parsed
                table_prefix, col_key = _resolve_repeat_path(spec)
                tokens.append(TokenOccurrence(
                    token=key,
                    table_prefix=table_prefix,
                    column_key=col_key,
                    location={
                        "type": "paragraph",
                        "paragraph_index": p_idx,
                        "source_text": text[:200],
                    },
                    placeholder=m.group(0),
                    hint_type=hint_type,
                    hint_args=hint_args,
                    segments=spec["segments"],
                    repeat_root=spec["repeat_root"],
                    params=spec["params"],
                ))

        # Tables
        for t_idx, table in enumerate(doc.tables):
            region_tokens: List[str] = []
            header_texts: List[str] = []
            prefix_counts: Dict[str, int] = {}
            prefix_row_signatures: Dict[str, List[Tuple[str, ...]]] = {}

            for r_idx, row in enumerate(table.rows):
                row_prefix_tokens: Dict[str, List[str]] = {}
                for c_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    if not cell_text:
                        continue
                    text_lines.append(cell_text)

                    for m in _TOKEN_RE.finditer(cell_text):
                        parsed = _parse_placeholder_expr(m.group(1))
                        if not parsed:
                            continue
                        key, hint_type, hint_args, spec = parsed
                        table_prefix, col_key = _resolve_repeat_path(spec)
                        tok = m.group(0)
                        if tok not in region_tokens:
                            region_tokens.append(tok)
                        tokens.append(TokenOccurrence(
                            token=key,
                            table_prefix=table_prefix,
                            column_key=col_key,
                            location={
                                "type": "table",
                                "table_index": t_idx,
                                "row_index": r_idx,
                                "col_index": c_idx,
                                "source_text": cell_text[:200],
                            },
                            placeholder=m.group(0),
                            hint_type=hint_type,
                            hint_args=hint_args,
                            segments=spec["segments"],
                            repeat_root=spec["repeat_root"],
                            params=spec["params"],
                        ))
                        if table_prefix:
                            prefix_counts[table_prefix] = prefix_counts.get(table_prefix, 0) + 1
                            row_prefix_tokens.setdefault(table_prefix, []).append(tok)

                    if r_idx == 0:
                        header_texts.append(cell_text)

                for prefix, toks in row_prefix_tokens.items():
                    prefix_row_signatures.setdefault(prefix, []).append(tuple(dict.fromkeys(toks)))

            if not table.rows:
                continue

            rows_count = len(table.rows)
            cols_count = len(table.columns)
            # Find dominant prefix (most column tokens)
            dominant_prefix = max(prefix_counts, key=lambda k: prefix_counts[k]) if prefix_counts else None
            loop_toks = [
                t.placeholder or f"{{{{{t.token}}}}}"
                for t in tokens
                if dominant_prefix and (
                    t.table_prefix == dominant_prefix
                )
            ] if dominant_prefix else []

            table_regions.append(TableRegion(
                region_id=f"docx:table:{t_idx}",
                location={
                    "type": "docx_table",
                    "table_index": t_idx,
                    "rows": rows_count,
                    "cols": cols_count,
                },
                loop_tokens=loop_toks,
                loop_prefix=dominant_prefix,
                header_row=header_texts,
                template_row_index=1 if header_texts else 0,
                orientation="vertical",
            ))

        # Unclosed fences
        for key, open_pos in open_fences.items():
            fence_blocks.append(FenceBlock(key=key, open_position=open_pos))

        title, version = _extract_title_version(text_lines[:20])
        scalar_keys, table_prefixes = _aggregate_keys(tokens, table_regions)

        return RawLayout(
            format="docx",
            title=title,
            version=version,
            tokens=tokens,
            table_regions=table_regions,
            fence_blocks=fence_blocks,
            scalar_keys=scalar_keys,
            table_prefixes=table_prefixes,
            schema_roots=_build_schema_roots(tokens, table_regions),
            text_lines=text_lines[:500],
        )

    # ------------------------------------------------------------------
    # Plain text
    # ------------------------------------------------------------------

    def _parse_text(self, content: bytes, filename: str) -> RawLayout:
        try:
            text = content.decode("utf-8")
        except UnicodeDecodeError:
            text = content.decode("latin-1", errors="replace")

        tokens: List[TokenOccurrence] = []
        fence_blocks: List[FenceBlock] = []
        open_fences: Dict[str, Dict[str, Any]] = {}
        text_lines: List[str] = []

        for line_idx, raw_line in enumerate(text.splitlines()):
            line = raw_line.strip()
            if not line:
                continue
            text_lines.append(line)

            for m in _FENCE_OPEN_RE.finditer(line):
                key = m.group(1)
                open_fences[key] = {"line_index": line_idx, "context": line[:100]}

            for m in _FENCE_CLOSE_RE.finditer(line):
                key = m.group(1)
                open_pos = open_fences.pop(key, None)
                fence_blocks.append(FenceBlock(
                    key=key,
                    open_position=open_pos or {},
                    close_position={"line_index": line_idx, "context": line[:100]},
                ))

            for m in _TOKEN_RE.finditer(line):
                parsed = _parse_placeholder_expr(m.group(1))
                if not parsed:
                    continue
                key, hint_type, hint_args, spec = parsed
                table_prefix, col_key = _resolve_repeat_path(spec)
                tokens.append(TokenOccurrence(
                    token=key,
                    table_prefix=table_prefix,
                    column_key=col_key,
                    location={
                        "type": "line",
                        "line_index": line_idx,
                        "source_text": line[:200],
                    },
                    placeholder=m.group(0),
                    hint_type=hint_type,
                    hint_args=hint_args,
                    segments=spec["segments"],
                    repeat_root=spec["repeat_root"],
                    params=spec["params"],
                ))

        for key, open_pos in open_fences.items():
            fence_blocks.append(FenceBlock(key=key, open_position=open_pos))

        # Table regions from fence blocks
        table_regions: List[TableRegion] = []
        for fb in fence_blocks:
            if fb.close_position is not None:
                region_tokens = [
                    t.placeholder or f"{{{{{t.token}}}}}"
                    for t in tokens
                    if t.table_prefix == fb.key or t.token == fb.key or t.token.startswith(f"{fb.key}.")
                ]
                table_regions.append(TableRegion(
                    region_id=f"text:fence:{fb.key}",
                    location={
                        "type": "fence",
                        "key": fb.key,
                        "open": fb.open_position,
                        "close": fb.close_position,
                    },
                    loop_tokens=list(dict.fromkeys(region_tokens)),
                    loop_prefix=fb.key,
                    orientation="vertical",
                ))

        # Also add inline marker regions for explicit repeats or fenced blocks.
        prefix_tokens: Dict[str, List[str]] = {}
        prefix_line_signatures: Dict[str, List[Tuple[str, ...]]] = {}
        for line in text_lines:
            line_prefixes: Dict[str, List[str]] = {}
            for m in _TOKEN_RE.finditer(line):
                parsed = _parse_placeholder_expr(m.group(1))
                if not parsed:
                    continue
                _, _, _, spec = parsed
                tp, _ = _resolve_repeat_path(spec)
                if tp:
                    line_prefixes.setdefault(tp, []).append(m.group(0))
            for prefix, toks in line_prefixes.items():
                prefix_line_signatures.setdefault(prefix, []).append(tuple(dict.fromkeys(toks)))
        for t in tokens:
            if t.table_prefix:
                prefix_tokens.setdefault(t.table_prefix, []).append(t.placeholder or f"{{{{{t.token}}}}}")
        for prefix, toks in prefix_tokens.items():
            signatures = prefix_line_signatures.get(prefix, [])
            repeated_signature = len({sig for sig in signatures if sig}) < len(signatures)
            explicit_repeat = any(
                bool(segment.get("repeat"))
                for token in tokens
                if token.table_prefix == prefix
                for segment in token.segments
            )
            if repeated_signature or explicit_repeat or any(fb.key == prefix for fb in fence_blocks):
                table_regions.append(TableRegion(
                    region_id=f"text:marker:{prefix}",
                    location={"type": "inline", "key": prefix},
                    loop_tokens=list(dict.fromkeys(toks)),
                    loop_prefix=prefix,
                    orientation="vertical",
                ))

        title, version = _extract_title_version(text_lines[:10])
        scalar_keys, table_prefixes = _aggregate_keys(tokens, table_regions)

        return RawLayout(
            format="text",
            title=title,
            version=version,
            tokens=tokens,
            table_regions=table_regions,
            fence_blocks=fence_blocks,
            scalar_keys=scalar_keys,
            table_prefixes=table_prefixes,
            schema_roots=_build_schema_roots(tokens, table_regions),
            text_lines=text_lines[:500],
        )


# ---------------------------------------------------------------------------
# Private helpers
# ---------------------------------------------------------------------------


def _ext(filename: str) -> str:
    return filename.rsplit(".", 1)[-1].strip().lower() if "." in filename else ""


def _split_dotted(key: str) -> Tuple[Optional[str], Optional[str]]:
    """Split a dotted key like 'items.qty' → ('items', 'qty').

    Non-dotted keys return ``(None, None)``.
    Only the *first* dot is used as separator; deeper nesting is not supported.
    """
    if "." in key:
        parts = key.split(".", 1)
        return parts[0], parts[1]
    return None, None


def _resolve_repeat_path(spec: Dict[str, Any]) -> Tuple[Optional[str], Optional[str]]:
    segments = spec.get("segments") or []
    if not segments:
        return None, None
    seen_repeat = False
    repeat_path_parts: List[str] = []
    tail: List[str] = []
    for segment in segments:
        key = str(segment.get("key") or "").strip()
        if not key:
            continue
        if not seen_repeat:
            repeat_path_parts.append(key)
        if bool(segment.get("repeat")) and not seen_repeat:
            seen_repeat = True
            continue
        if seen_repeat:
            tail.append(key)
    if not seen_repeat:
        return None, None
    return ".".join(repeat_path_parts), ".".join(tail) if tail else None


def _extract_title_version(texts: List[str]) -> Tuple[Optional[str], Optional[str]]:
    title: Optional[str] = None
    version: Optional[str] = None
    for raw in texts:
        line = raw.strip()
        if not line:
            continue
        if title is None and len(line) > 3 and not _TOKEN_RE.search(line):
            title = line
        if version is None:
            m = _VERSION_RE.search(line)
            if m:
                version = m.group(1)
        if title and version:
            break
    return title, version


def _aggregate_keys(
    tokens: List[TokenOccurrence],
    table_regions: List[TableRegion],
) -> Tuple[List[str], List[str]]:
    """Return (scalar_keys, table_prefixes) from token list, deduplicated, ordered."""
    scalars: List[str] = []
    prefixes: List[str] = []
    seen_s: set = set()
    seen_p: set = set()
    table_prefix_set = {
        region.loop_prefix
        for region in table_regions
        if region.loop_prefix
    }
    for t in tokens:
        if t.table_prefix and t.table_prefix in table_prefix_set:
            if t.table_prefix not in seen_p:
                seen_p.add(t.table_prefix)
                prefixes.append(t.table_prefix)
        else:
            if t.token not in seen_s:
                seen_s.add(t.token)
                scalars.append(t.token)
    return scalars, prefixes


def _parse_placeholder_expr(expr: str) -> Optional[Tuple[str, Optional[str], Optional[str], Dict[str, Any]]]:
    raw = expr.strip()
    if not raw or raw.startswith("#") or raw.startswith("/"):
        return None
    legacy_hint_match = re.fullmatch(
        r"(?P<path>[A-Za-z0-9_.\-]+):(?P<hint>[A-Za-z0-9_\-]+)(?:\((?P<args>.*)\))?",
        raw,
    )
    parse_target = raw
    legacy_hint_type = None
    legacy_hint_args = None
    if legacy_hint_match:
        parse_target = str(legacy_hint_match.group("path") or "").strip()
        legacy_hint_type = str(legacy_hint_match.group("hint") or "").strip().lower() or None
        legacy_hint_args = str(legacy_hint_match.group("args") or "").strip() or None

    spec = _parse_contract_expr(parse_target)
    if spec is None:
        return None
    key = spec["path"]
    params = spec["params"]
    leaf_params = spec.get("leaf_params") or {}
    hint_type = legacy_hint_type
    hint_args = legacy_hint_args
    if "type" in leaf_params:
        hint_type = str(leaf_params["type"]).strip().lower() or None
        if "type_args" in leaf_params and not hint_args:
            hint_args = str(leaf_params["type_args"]).strip() or None
        if "min" in leaf_params or "max" in leaf_params:
            extras = []
            if "min" in leaf_params:
                extras.append(f"min={leaf_params['min']}")
            if "max" in leaf_params:
                extras.append(f"max={leaf_params['max']}")
            hint_args = ", ".join(extras) if extras else None
    return key, hint_type, hint_args, spec


def _parse_contract_expr(expr: str) -> Optional[Dict[str, Any]]:
    parts = _split_path_segments(expr)
    if not parts:
        return None

    segments: List[PathSegment] = []
    path_parts: List[str] = []
    repeat_root = None
    merged_params: Dict[str, Any] = {}

    for idx, part in enumerate(parts):
        segment = part.strip()
        if not segment:
            return None
        repeat = segment.endswith("[]")
        if repeat:
            segment = segment[:-2].rstrip()

        params: Dict[str, Any] = {}
        if "(" in segment:
            head, tail = segment.split("(", 1)
            if not tail.endswith(")"):
                return None
            segment = head.strip()
            params = _parse_params(tail[:-1])
        if not segment or not re.fullmatch(r"[A-Za-z0-9_\-]+", segment):
            return None
        path_parts.append(segment)
        if idx == len(parts) - 1:
            merged_params.update(params)
        segments.append(PathSegment(key=segment, params=params, repeat=repeat))
        if repeat and repeat_root is None:
            repeat_root = segment

    path = ".".join(path_parts)
    segment_dicts = [asdict(segment) for segment in segments]
    repeat_params: Dict[str, Any] = {}
    for segment in segments:
        if segment.repeat:
            repeat_params = dict(segment.params)
            break
    leaf_params = dict(segments[-1].params) if segments else {}
    return {
        "path": path,
        "segments": segment_dicts,
        "repeat_root": repeat_root,
        "params": merged_params,
        "leaf_params": leaf_params,
        "repeat_params": repeat_params,
    }


def _split_path_segments(expr: str) -> List[str]:
    parts: List[str] = []
    current: List[str] = []
    depth = 0
    for ch in expr:
        if ch == "." and depth == 0:
            parts.append("".join(current))
            current = []
            continue
        if ch == "(":
            depth += 1
        elif ch == ")":
            depth = max(0, depth - 1)
        current.append(ch)
    if current:
        parts.append("".join(current))
    return parts


def _parse_params(raw: str) -> Dict[str, Any]:
    params: Dict[str, Any] = {}
    if not raw.strip():
        return params
    for chunk in _split_args(raw):
        if not chunk:
            continue
        if "=" not in chunk:
            chunk = chunk.strip()
            shorthand = _parse_type_shorthand(chunk)
            if shorthand is not None:
                params["type"] = shorthand["type"]
                if shorthand.get("args"):
                    params["type_args"] = shorthand["args"]
            else:
                params[chunk] = True
            continue
        key, value = chunk.split("=", 1)
        params[key.strip()] = _coerce_param_value(value.strip())
    return params


def _parse_type_shorthand(raw: str) -> Optional[Dict[str, str]]:
    value = raw.strip()
    if not value:
        return None
    match = re.fullmatch(r"(?P<type>[A-Za-z0-9_\-]+)(?:\((?P<args>.*)\))?", value)
    if not match:
        return None
    type_name = str(match.group("type") or "").strip()
    if not type_name:
        return None
    return {
        "type": type_name,
        "args": str(match.group("args") or "").strip(),
    }


def _split_args(raw: str) -> List[str]:
    parts: List[str] = []
    current: List[str] = []
    depth_paren = 0
    depth_bracket = 0
    depth_brace = 0
    quote: Optional[str] = None
    for ch in raw:
        if quote:
            current.append(ch)
            if ch == quote:
                quote = None
            continue
        if ch in {"'", '"'}:
            quote = ch
            current.append(ch)
            continue
        if ch == "(":
            depth_paren += 1
        elif ch == ")":
            depth_paren = max(0, depth_paren - 1)
        elif ch == "[":
            depth_bracket += 1
        elif ch == "]":
            depth_bracket = max(0, depth_bracket - 1)
        elif ch == "{":
            depth_brace += 1
        elif ch == "}":
            depth_brace = max(0, depth_brace - 1)
        if ch == "," and depth_paren == 0 and depth_bracket == 0 and depth_brace == 0:
            parts.append("".join(current).strip())
            current = []
            continue
        current.append(ch)
    if current:
        parts.append("".join(current).strip())
    return parts


def _build_schema_roots(tokens: List[TokenOccurrence], table_regions: List[TableRegion]) -> List[SchemaNode]:
    roots: Dict[str, SchemaNode] = {}
    for token in tokens:
        _merge_token_into_tree(roots, token)
    _attach_anchor_hints(roots, table_regions)
    return list(roots.values())


def _merge_token_into_tree(roots: Dict[str, SchemaNode], token: TokenOccurrence) -> None:
    current_map = roots
    parent_node: Optional[SchemaNode] = None
    path_parts: List[str] = []
    current_node: Optional[SchemaNode] = None
    segments = token.segments or _segments_from_token(token)
    for index, segment in enumerate(segments):
        key = str(segment.get("key") or "").strip()
        if not key:
            return
        path_parts.append(key)
        path = ".".join(path_parts)
        node = current_map.get(key)
        if node is None:
            node = SchemaNode(
                key=key,
                path=path,
                params=dict(segment.get("params") or {}),
                repeat=bool(segment.get("repeat")),
            )
            current_map[key] = node
            if parent_node is not None:
                parent_node.children = list(current_map.values())
        else:
            node.params.update(segment.get("params") or {})
            node.repeat = node.repeat or bool(segment.get("repeat"))
        if index == len(segments) - 1:
            node.placeholder = node.placeholder or token.placeholder
            if not node.location:
                node.location = dict(token.location)
            node.hint_type = node.hint_type or token.hint_type
            node.hint_args = node.hint_args or token.hint_args
        current_node = node
        parent_node = node
        current_map = {child.key: child for child in node.children}
    if current_node is not None and current_node.placeholder is None:
        current_node.placeholder = token.placeholder


def _segments_from_token(token: TokenOccurrence) -> List[Dict[str, Any]]:
    if not token.token:
        return []
    parts = token.token.split(".")
    repeat_parts = token.table_prefix.split(".") if token.table_prefix else []
    segments: List[Dict[str, Any]] = []
    for index, part in enumerate(parts):
        current_path = parts[: index + 1]
        segments.append(
            {
                "key": part,
                "params": {},
                "repeat": bool(repeat_parts) and current_path == repeat_parts,
            }
        )
    return segments


def _attach_anchor_hints(roots: Dict[str, SchemaNode], table_regions: List[TableRegion]) -> None:
    nodes_by_path: Dict[str, SchemaNode] = {}

    def _walk(node: SchemaNode) -> None:
        nodes_by_path[node.path] = node
        for child in node.children:
            _walk(child)

    for root in roots.values():
        _walk(root)

    for region in table_regions:
        if not region.loop_prefix:
            continue
        node = nodes_by_path.get(region.loop_prefix)
        if node is None:
            continue
        node.anchors.append(
            {
                "strategy": "marker" if region.loop_tokens else "structural",
                "loop_tokens": list(region.loop_tokens),
                "marker_columns": dict(region.marker_columns),
                "header_row": list(region.header_row or []),
                "location": dict(region.location),
                "orientation": region.orientation,
                "template_row_index": region.template_row_index,
            }
        )


def _coerce_param_value(raw: str) -> Any:
    import ast

    text = raw.strip()
    if not text:
        return ""
    try:
        return ast.literal_eval(text)
    except Exception:
        lowered = text.lower()
        if lowered in {"true", "false"}:
            return lowered == "true"
        return text.strip("'\"")
