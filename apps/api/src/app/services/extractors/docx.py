"""DOCX extractor with style-aware markdown conversion."""
from __future__ import annotations

import re
from io import BytesIO
from typing import List, Set

from app.services.extractors.base import BaseExtractor, ExtractResult


class DocxExtractor(BaseExtractor):
    """Extracts text from DOCX files preserving headings, lists, and tables."""

    @property
    def extensions(self) -> Set[str]:
        return {"docx"}

    @property
    def kind(self) -> str:
        return "docx"

    def extract(self, data: bytes, filename: str) -> ExtractResult:
        warnings: List[str] = []
        text = ""
        try:
            from docx import Document  # type: ignore

            doc = Document(BytesIO(data))
            parts: List[str] = []

            # Process paragraphs with style awareness
            for p in doc.paragraphs:
                if not p.text:
                    continue

                style_name = p.style.name if p.style else ""
                p_text = p.text.strip()

                # Map styles to Markdown
                if style_name.startswith("Heading 1"):
                    parts.append(f"\n# {p_text}\n")
                elif style_name.startswith("Heading 2"):
                    parts.append(f"\n## {p_text}\n")
                elif style_name.startswith("Heading 3"):
                    parts.append(f"\n### {p_text}\n")
                elif style_name.startswith("List"):
                    parts.append(f"- {p_text}")
                else:
                    parts.append(p_text)

            # Process tables (simple conversion to text/tabs)
            for t in doc.tables:
                parts.append("\n")
                for row in t.rows:
                    cells = [c.text.strip().replace("\n", " ") for c in row.cells]
                    if any(cells):
                        parts.append("| " + " | ".join(cells) + " |")
                parts.append("\n")

            text = "\n".join(parts).strip()
            text = re.sub(r"\n{3,}", "\n\n", text)

            if not text:
                warnings.append("DOCX parsed but no visible text was found.")
        except Exception as e:
            warnings.append(f"DOCX extraction failed: {e!r}")
            text = ""

        return ExtractResult(text=text, kind="docx", meta={}, warnings=warnings)
