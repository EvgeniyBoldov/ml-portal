"""
collection.template.fill — fill a template row and produce the final file artifact.

Supports Excel, Word, and plain text templates via contract-aware filling.
The result is a stored downloadable file artifact and should be treated as the
final output of the fill operation.
"""
from __future__ import annotations

import io
import re
from pathlib import PurePosixPath
from typing import Any, ClassVar, Dict

from app.agents.context import ToolContext, ToolResult
from app.agents.handlers.versioned_tool import VersionedTool, register_tool, tool_version
from app.core.db import get_session_factory
from app.core.logging import get_logger
from app.models.collection import CollectionType
from app.services.artifact_writer import ArtifactWriter
from app.services.collection.row_service import CollectionRowService
from app.services.collection.template_contract import TemplateContract
from app.services.collection.template_fill_engine import TemplateFillEngine
from app.services.collection.template_layout_parser import _parse_placeholder_expr
from app.services.collection_service import CollectionService
from app.services.file_delivery_service import FileDeliveryService
from app.adapters.s3_client import s3_manager

logger = get_logger(__name__)

_PLACEHOLDER_RE = re.compile(r"\{\{([^{}]+)\}\}")

_INPUT_SCHEMA_V1 = {
    "type": "object",
    "properties": {
        "collection_slug": {
            "type": "string",
            "description": "Slug of the template collection",
        },
        "row_id": {
            "type": "string",
            "description": "UUID of the template row",
        },
        "values": {
            "type": "object",
            "description": "Values payload for the selected template row. Keys, nesting, and types must match the stored fill schema.",
        },
        "filename": {
            "type": "string",
            "description": "Optional output filename. Its extension must match the template format.",
        },
    },
    "required": ["collection_slug", "row_id", "values"],
}

_OUTPUT_SCHEMA_V1 = {
    "type": "object",
    "properties": {
        "artifact_id": {"type": "string"},
        "file_name": {"type": "string"},
        "content_type": {"type": "string"},
        "size_bytes": {"type": "integer"},
        "format": {"type": "string"},
        "filled_placeholders": {"type": "integer"},
        "missing_placeholders": {"type": "array", "items": {"type": "string"}},
    },
}


def _fill_text(content: bytes, values: Dict[str, str]) -> bytes:
    text = content.decode("utf-8")
    text, _ = _substitute_placeholders(text, values)
    return text.encode("utf-8")


def _fill_excel(content: bytes, values: Dict[str, str]) -> bytes:
    try:
        import openpyxl
    except ImportError:
        raise RuntimeError("openpyxl is not installed; cannot fill Excel templates")

    wb = openpyxl.load_workbook(io.BytesIO(content))
    filled = set()
    for sheet in wb.worksheets:
        for row in sheet.iter_rows():
            for cell in row:
                if cell.value and isinstance(cell.value, str):
                    new_val, keys = _substitute_placeholders(cell.value, values)
                    if new_val != cell.value:
                        cell.value = new_val
                        filled.update(keys)
    out = io.BytesIO()
    wb.save(out)
    wb.close()
    return out.getvalue()


def _fill_word(content: bytes, values: Dict[str, str]) -> bytes:
    try:
        import docx
    except ImportError:
        raise RuntimeError("python-docx is not installed; cannot fill Word templates")

    doc = docx.Document(io.BytesIO(content))
    filled = set()

    for para in doc.paragraphs:
        if para.text:
            new_text, keys = _substitute_placeholders(para.text, values)
            if new_text != para.text:
                para.clear()
                para.add_run(new_text)
                filled.update(keys)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    new_text, keys = _substitute_placeholders(cell.text, values)
                    if new_text != cell.text:
                        cell.paragraphs[0].clear()
                        cell.paragraphs[0].add_run(new_text)
                        filled.update(keys)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _substitute_placeholders(text: str, values: Dict[str, str]) -> tuple[str, set[str]]:
    keys_used = set()

    def replacer(match: Any) -> str:
        parsed = _parse_placeholder_expr(match.group(1))
        if not parsed:
            return match.group(0)
        key, _, _, _ = parsed
        if key in values:
            keys_used.add(key)
            return str(values[key])
        return match.group(0)

    result = _PLACEHOLDER_RE.sub(replacer, text)
    return result, keys_used
@register_tool
class TemplateFillTool(VersionedTool):
    """Fill a template row and return the final generated file artifact."""

    tool_slug: ClassVar[str] = "collection.template.fill"
    domains: ClassVar[list] = ["collection.template"]
    name: ClassVar[str] = "Fill Template"
    description: ClassVar[str] = (
        "Fill a template row with provided values and create the final downloadable file artifact. "
        "Returns the resulting chat artifact_id. Use it as the file reference. "
        "The values object must match the stored fill schema for the selected template row."
    )

    @tool_version(
        version="1.0.0",
        input_schema=_INPUT_SCHEMA_V1,
        output_schema=_OUTPUT_SCHEMA_V1,
        description="Fill template and return file",
    )
    async def v1_0_0(self, ctx: ToolContext, args: Dict[str, Any]) -> ToolResult:
        log = ctx.tool_notes("collection.template.fill")

        collection_slug = str(args.get("collection_slug") or "").strip()
        # Existing persisted task payloads may still contain collection_id.
        # It is a compatibility input only; all published contracts use slug.
        legacy_collection_id = str(args.get("collection_id") or "").strip()
        row_id = str(args.get("row_id") or "").strip()
        values = args.get("values") or {}

        if not (collection_slug or legacy_collection_id) or not row_id:
            log.error("Missing collection_slug or row_id")
            return ToolResult.fail(
                "Missing 'collection_slug' or 'row_id' argument",
                logs=log.entries_dict(),
            )
        if not isinstance(values, dict):
            log.error("Invalid values type", type=type(values).__name__)
            return ToolResult.fail("'values' must be an object/dict", logs=log.entries_dict())

        try:
            import uuid
            session_factory = get_session_factory()
            async with session_factory() as session:
                service = CollectionService(session)
                if collection_slug:
                    collection = await service.get_by_slug(collection_slug)
                else:
                    try:
                        collection = await service.get_by_id(uuid.UUID(legacy_collection_id))
                    except ValueError:
                        collection = None

                if not collection:
                    return ToolResult.fail(
                        f"Collection '{collection_slug or legacy_collection_id}' not found",
                        logs=log.entries_dict(),
                    )
                if collection.collection_type != CollectionType.TEMPLATE.value:
                    return ToolResult.fail(
                        f"Collection '{collection_slug or legacy_collection_id}' is not a template collection",
                        logs=log.entries_dict(),
                    )

                rid = uuid.UUID(row_id)
                row_service = CollectionRowService(session)
                row = await row_service.get_row_by_id(collection, rid)
                if not row:
                    return ToolResult.fail(
                        f"Template row '{row_id}' not found",
                        logs=log.entries_dict(),
                    )

                if str(row.get("status") or "").lower() != "ready":
                    return ToolResult.fail(
                        "Template is not ready. It must finish analysis, be approved, and be vectorized before filling.",
                        logs=log.entries_dict(),
                    )
                file_meta = row.get("file") or {}
                s3_key = file_meta.get("s3_key")
                bucket = file_meta.get("bucket")
                filename = file_meta.get("filename") or "template"
                if not s3_key or not bucket:
                    return ToolResult.fail(
                        "Template file metadata is incomplete (missing s3_key or bucket)",
                        logs=log.entries_dict(),
                    )

                requested_filename = str(args.get("filename") or "").strip()
                output_filename = requested_filename or f"filled_{filename}"
                source_ext = PurePosixPath(str(filename)).suffix.lower()
                output_ext = PurePosixPath(output_filename).suffix.lower()
                if source_ext not in {".xlsx", ".xlsm"}:
                    return ToolResult.fail("Only .xlsx and .xlsm templates are supported", logs=log.entries_dict())
                if requested_filename and source_ext and output_ext != source_ext:
                    return ToolResult.fail(
                        "Output filename extension must match the template format",
                        logs=log.entries_dict(),
                    )

                # Download template from S3
                content = await s3_manager.get_object(bucket, s3_key)
                if content is None:
                    return ToolResult.fail(
                        f"Failed to load template file from storage: {bucket}/{s3_key}",
                        logs=log.entries_dict(),
                    )

                # Determine format for response
                ext = ""
                if "." in filename:
                    ext = filename.rsplit(".", 1)[-1].strip().lower()
                if ext in {"xlsx", "xlsm"}:
                    fmt = "excel"
                else:
                    return ToolResult.fail("Only .xlsx and .xlsm templates are supported", logs=log.entries_dict())

                # Load contract
                raw_schema = row.get("template_schema") or {}
                contract = TemplateContract.from_jsonb(raw_schema)

                if not contract.fields:
                    return ToolResult.fail(
                        "Template schema is not ready for this row. Fill is available only after schema analysis completes.",
                        logs=log.entries_dict(),
                    )

                validation = contract.validate_generated_values(values)
                if not validation.ok:
                    log.warning(
                        "template_fill_validation_failed",
                        error_count=len(validation.error_details),
                        warning_count=len(validation.warning_details),
                    )
                    validation_errors = [
                        issue.model_dump(mode="json", exclude_none=True)
                        for issue in validation.error_details
                    ]
                    validation_warnings = [
                        issue.model_dump(mode="json", exclude_none=True)
                        for issue in validation.warning_details
                    ]
                    return ToolResult.fail(
                        (
                            "Provided values JSON does not match the template schema. "
                            "Fix the invalid fields listed in metadata.validation_errors and call "
                            "'collection.template.fill' again with corrected 'values'."
                        ),
                        logs=log.entries_dict(),
                        validation_summary={
                            "valid": False,
                            "error_count": len(validation_errors),
                            "warning_count": len(validation_warnings),
                        },
                        validation_errors=validation_errors,
                        validation_warnings=validation_warnings,
                        retry_hint=(
                            "Correct each field listed in metadata.validation_errors, keep the same nested JSON shape, "
                            "and repeat the tool call with updated values."
                        ),
                    )

                engine = TemplateFillEngine(contract)
                result = engine.fill(content, values, filename, assume_valid=True)
                if not result.success:
                    return ToolResult.fail(
                        f"Failed to fill template: {result.error}",
                        logs=log.entries_dict(),
                    )
                filled_bytes = result.content
                filled_keys = set(result.filled_scalars + result.filled_tables)
                missing = list(set(result.missing_scalars + result.missing_tables))

                # Store generated attachment
                chat_id = str(ctx.chat_id) if ctx.chat_id else None
                owner_id = str(ctx.user_id or "")
                if not chat_id:
                    return ToolResult.fail(
                        "Template filling requires a chat context.",
                        logs=log.entries_dict(),
                    )
                if not owner_id:
                    return ToolResult.fail(
                        "Tool context missing user_id; cannot store generated file",
                        logs=log.entries_dict(),
                    )

                artifact = await ArtifactWriter(session).write(
                    chat_id=chat_id,
                    owner_id=owner_id,
                    filename=output_filename,
                    content=filled_bytes,
                    content_type=file_meta.get("content_type") or "application/octet-stream",
                    metadata={"format": fmt, "template_filename": filename},
                )
                await session.commit()

                return ToolResult.ok(
                    data={
                        "artifact_id": artifact.artifact_id,
                        "file_name": artifact.file_name,
                        "content_type": file_meta.get("content_type") or "application/octet-stream",
                        "size_bytes": artifact.size_bytes,
                        "format": fmt,
                        "filled_placeholders": len(filled_keys),
                        "missing_placeholders": missing,
                    },
                    message=(
                        f"Final downloadable file created from template '{filename}' "
                        f"({fmt}, {len(filled_keys)} placeholders)."
                    ),
                    logs=log.entries_dict(),
                )
        except Exception as exc:
            logger.error("collection.template.fill failed: %s", exc, exc_info=True)
            log.error("collection.template.fill failed", error=str(exc))
            return ToolResult.fail(f"Failed to fill template: {exc}", logs=log.entries_dict())
