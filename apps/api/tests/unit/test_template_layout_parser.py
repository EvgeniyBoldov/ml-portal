"""Unit tests for TemplateLayoutParser (S1)."""
from __future__ import annotations

import pytest

from app.services.collection.template_layout_parser import (
    TemplateLayoutParser,
    RawLayout,
    _split_dotted,
    _extract_title_version,
)


@pytest.fixture
def parser() -> TemplateLayoutParser:
    return TemplateLayoutParser()


# ---------------------------------------------------------------------------
# _split_dotted helper
# ---------------------------------------------------------------------------


def test_split_dotted_simple():
    assert _split_dotted("name") == (None, None)


def test_split_dotted_dotted():
    assert _split_dotted("items.qty") == ("items", "qty")


def test_split_dotted_deeper():
    # Only first dot is used
    assert _split_dotted("a.b.c") == ("a", "b.c")


# ---------------------------------------------------------------------------
# _extract_title_version helper
# ---------------------------------------------------------------------------


def test_extract_title_only():
    title, ver = _extract_title_version(["Заявка на оборудование", "подразделение: ИТ"])
    assert title == "Заявка на оборудование"
    assert ver is None


def test_extract_version():
    _, ver = _extract_title_version(["Форма Version 3.1", "поле"])
    assert ver == "3.1"


def test_extract_version_russian():
    _, ver = _extract_title_version(["Шаблон версия: 2.0"])
    assert ver == "2.0"


def test_extract_skips_token_lines():
    title, _ = _extract_title_version(["{{skip_me}}", "Real Title"])
    assert title == "Real Title"


# ---------------------------------------------------------------------------
# Text parsing
# ---------------------------------------------------------------------------


def test_parse_text_scalars(parser: TemplateLayoutParser):
    content = b"Hello {{name}}, your code is {{code}}."
    layout = parser.parse(content, "greeting.txt")
    assert layout.format == "text"
    keys = {t.token for t in layout.tokens}
    assert "name" in keys
    assert "code" in keys
    assert "name" in layout.scalar_keys
    assert "code" in layout.scalar_keys
    assert layout.table_prefixes == []


def test_parse_text_dotted_tokens(parser: TemplateLayoutParser):
    content = b"{{items.name}} {{items.qty}} {{items.price}}"
    layout = parser.parse(content, "form.txt")
    assert "items" in layout.table_prefixes
    assert layout.scalar_keys == []


def test_parse_text_dotted_object_tokens(parser: TemplateLayoutParser):
    content = b"{{author.name}} {{author.email}}"
    layout = parser.parse(content, "form.txt")
    assert layout.table_prefixes == []
    assert "author.name" in layout.scalar_keys
    assert "author.email" in layout.scalar_keys


def test_parse_text_table_region_from_dotted(parser: TemplateLayoutParser):
    content = b"{{items.name}} {{items.qty}}"
    layout = parser.parse(content, "form.txt")
    assert len(layout.table_regions) == 1
    region = layout.table_regions[0]
    assert region.loop_prefix == "items"
    assert "{{items.name}}" in region.loop_tokens
    assert "{{items.qty}}" in region.loop_tokens


def test_parse_text_fence_blocks(parser: TemplateLayoutParser):
    content = b"{{#rows}}\n{{rows.name}} | {{rows.amount}}\n{{/rows}}"
    layout = parser.parse(content, "report.txt")
    assert len(layout.fence_blocks) == 1
    assert layout.fence_blocks[0].key == "rows"
    assert layout.fence_blocks[0].close_position is not None


def test_parse_text_unclosed_fence(parser: TemplateLayoutParser):
    content = b"{{#items}}\n{{items.name}}"
    layout = parser.parse(content, "broken.txt")
    # unclosed fence is still recorded
    assert any(fb.key == "items" for fb in layout.fence_blocks)


def test_parse_text_title_extracted(parser: TemplateLayoutParser):
    content = b"Order Form v1.2\n{{customer}}\n{{date}}"
    layout = parser.parse(content, "order.txt")
    assert layout.title == "Order Form v1.2"
    assert layout.version == "1.2"


def test_parse_text_no_tokens(parser: TemplateLayoutParser):
    content = b"Static content only."
    layout = parser.parse(content, "static.txt")
    assert layout.tokens == []
    assert layout.scalar_keys == []
    assert layout.table_regions == []


def test_parse_text_mixed_scalar_and_table(parser: TemplateLayoutParser):
    content = b"Org: {{org}}\n{{items.name}} {{items.qty}}"
    layout = parser.parse(content, "mixed.txt")
    assert "org" in layout.scalar_keys
    assert "items" in layout.table_prefixes


def test_parse_text_single_column_table_has_region(parser: TemplateLayoutParser):
    # A single-column dotted token must still produce a table region so that
    # table_prefixes and table_regions stay in sync.
    content = b"{{items.name}}"
    layout = parser.parse(content, "single.txt")
    assert "items" in layout.table_prefixes
    assert any(r.loop_prefix == "items" for r in layout.table_regions)


# ---------------------------------------------------------------------------
# Excel parsing
# ---------------------------------------------------------------------------


@pytest.fixture
def simple_excel_bytes():
    pytest.importorskip("openpyxl")
    import io
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    ws["A1"] = "Applicant"
    ws["B1"] = "{{applicant_name}}"
    ws["A2"] = "Code"
    ws["B2"] = "{{code}}"
    buf = io.BytesIO()
    wb.save(buf)
    wb.close()
    return buf.getvalue()


@pytest.fixture
def table_excel_bytes():
    pytest.importorskip("openpyxl")
    import io
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Заявка"
    # Header row
    ws["A1"] = "Наименование"
    ws["B1"] = "Кол-во"
    ws["C1"] = "Цена"
    # Marker row with dotted tokens
    ws["A2"] = "{{items.name}}"
    ws["B2"] = "{{items.qty}}"
    ws["C2"] = "{{items.price}}"
    buf = io.BytesIO()
    wb.save(buf)
    wb.close()
    return buf.getvalue()


@pytest.fixture
def two_tables_excel_bytes():
    pytest.importorskip("openpyxl")
    import io
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Form"
    # Table 1: items
    ws["A1"] = "Наименование"
    ws["B1"] = "Кол-во"
    ws["A2"] = "{{items.name}}"
    ws["B2"] = "{{items.qty}}"
    # Gap
    # Table 2: contacts
    ws["A5"] = "ФИО"
    ws["B5"] = "Телефон"
    ws["A6"] = "{{contacts.name}}"
    ws["B6"] = "{{contacts.phone}}"
    buf = io.BytesIO()
    wb.save(buf)
    wb.close()
    return buf.getvalue()


def test_parse_excel_scalar_tokens(parser, simple_excel_bytes):
    layout = parser.parse(simple_excel_bytes, "form.xlsx")
    assert layout.format == "excel"
    assert "applicant_name" in layout.scalar_keys
    assert "code" in layout.scalar_keys


def test_parse_excel_sheet_names(parser, simple_excel_bytes):
    layout = parser.parse(simple_excel_bytes, "form.xlsx")
    assert "Sheet1" in layout.sheets


def test_parse_excel_marker_table_region(parser, table_excel_bytes):
    layout = parser.parse(table_excel_bytes, "request.xlsx")
    assert "items" in layout.table_prefixes
    assert len(layout.table_regions) >= 1
    region = next(r for r in layout.table_regions if r.loop_prefix == "items")
    assert "{{items.name}}" in region.loop_tokens
    assert "{{items.qty}}" in region.loop_tokens
    assert "{{items.price}}" in region.loop_tokens


def test_parse_excel_two_tables(parser, two_tables_excel_bytes):
    layout = parser.parse(two_tables_excel_bytes, "two.xlsx")
    prefixes = set(layout.table_prefixes)
    assert "items" in prefixes
    assert "contacts" not in prefixes
    assert "contacts.name" in layout.scalar_keys
    assert "contacts.phone" in layout.scalar_keys
    regions_by_prefix = {r.loop_prefix for r in layout.table_regions if r.loop_prefix}
    assert "items" in regions_by_prefix
    assert "contacts" not in regions_by_prefix


def test_parse_excel_token_location(parser, table_excel_bytes):
    layout = parser.parse(table_excel_bytes, "request.xlsx")
    tok = next(t for t in layout.tokens if t.token == "items.name")
    assert tok.location["sheet"] == "Заявка"
    assert "row" in tok.location
    assert "col" in tok.location


def test_parse_excel_formula_token(parser):
    pytest.importorskip("openpyxl")
    import io
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    ws["A1"] = "Field"
    ws["B1"] = '="{{name}}"'

    buf = io.BytesIO()
    wb.save(buf)
    wb.close()

    layout = parser.parse(buf.getvalue(), "formula.xlsx")

    assert "name" in layout.scalar_keys
    token = next(t for t in layout.tokens if t.token == "name")
    assert token.location["coordinate"] == "B1"


def test_parse_excel_typed_placeholder_scalar(parser):
    pytest.importorskip("openpyxl")
    import io
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    ws["A1"] = "{{author.tel:int(10)}}"

    buf = io.BytesIO()
    wb.save(buf)
    wb.close()

    layout = parser.parse(buf.getvalue(), "typed.xlsx")

    assert "author.tel" in layout.scalar_keys
    token = next(t for t in layout.tokens if t.token == "author.tel")
    assert token.hint_type == "int"
    assert token.hint_args == "10"
    assert token.placeholder == "{{author.tel:int(10)}}"


def test_parse_excel_repeated_row_detects_table(parser):
    pytest.importorskip("openpyxl")
    import io
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    ws["A1"] = "{{contacts.name}}"
    ws["B1"] = "{{contacts.phone}}"
    ws["A2"] = "{{contacts.name}}"
    ws["B2"] = "{{contacts.phone}}"

    buf = io.BytesIO()
    wb.save(buf)
    wb.close()

    layout = parser.parse(buf.getvalue(), "repeated.xlsx")

    assert "contacts" in layout.table_prefixes
    region = next(r for r in layout.table_regions if r.loop_prefix == "contacts")
    assert region.loop_tokens == ["{{contacts.name}}", "{{contacts.phone}}"]


# ---------------------------------------------------------------------------
# Docx parsing
# ---------------------------------------------------------------------------


@pytest.fixture
def simple_docx_bytes():
    pytest.importorskip("docx")
    import io
    import docx
    doc = docx.Document()
    doc.add_paragraph("Заявка на оборудование Version 1.0")
    doc.add_paragraph("Заявитель: {{applicant}}")
    doc.add_paragraph("Дата: {{date}}")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.fixture
def table_docx_bytes():
    pytest.importorskip("docx")
    import io
    import docx
    doc = docx.Document()
    doc.add_paragraph("Список позиций")
    table = doc.add_table(rows=2, cols=3)
    table.rows[0].cells[0].text = "Наименование"
    table.rows[0].cells[1].text = "Кол-во"
    table.rows[0].cells[2].text = "Цена"
    table.rows[1].cells[0].text = "{{items.name}}"
    table.rows[1].cells[1].text = "{{items.qty}}"
    table.rows[1].cells[2].text = "{{items.price}}"
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_parse_docx_scalar_tokens(parser, simple_docx_bytes):
    layout = parser.parse(simple_docx_bytes, "letter.docx")
    assert layout.format == "docx"
    assert "applicant" in layout.scalar_keys
    assert "date" in layout.scalar_keys


def test_parse_docx_title_version(parser, simple_docx_bytes):
    layout = parser.parse(simple_docx_bytes, "letter.docx")
    assert layout.title is not None
    assert "Заявка" in layout.title
    assert layout.version == "1.0"


def test_parse_docx_table_region(parser, table_docx_bytes):
    layout = parser.parse(table_docx_bytes, "items.docx")
    assert "items" in layout.table_prefixes
    region = next(r for r in layout.table_regions if r.loop_prefix == "items")
    assert "{{items.name}}" in region.loop_tokens


def test_parse_docx_table_header_row(parser, table_docx_bytes):
    layout = parser.parse(table_docx_bytes, "items.docx")
    region = next(r for r in layout.table_regions if r.loop_prefix == "items")
    assert region.header_row  # has header text extracted
    assert any("Наименование" in h for h in region.header_row)


def test_parse_docx_fence_block(parser):
    pytest.importorskip("docx")
    import io
    import docx
    doc = docx.Document()
    doc.add_paragraph("{{#lines}}")
    doc.add_paragraph("{{lines.text}}")
    doc.add_paragraph("{{/lines}}")
    buf = io.BytesIO()
    doc.save(buf)
    layout = parser.parse(buf.getvalue(), "fence.docx")
    assert any(fb.key == "lines" for fb in layout.fence_blocks)


# ---------------------------------------------------------------------------
# Edge cases
# ---------------------------------------------------------------------------


def test_parse_text_latin1_fallback(parser):
    content = "Привет {{name}}".encode("latin-1", errors="replace")
    layout = parser.parse(content, "test.txt")
    # Should not raise; tokens may or may not be found depending on encoding
    assert isinstance(layout, RawLayout)


def test_parse_unknown_extension_treated_as_text(parser):
    content = b"{{field1}} and {{field2}}"
    layout = parser.parse(content, "template.csv")
    assert layout.format == "text"
    assert "field1" in layout.scalar_keys


def test_excel_structural_fallback_no_markers(parser):
    pytest.importorskip("openpyxl")
    import io
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Data"
    ws["A1"] = "Name"
    ws["B1"] = "Value"
    ws["A2"] = "Alpha"
    ws["B2"] = "1"
    ws["A3"] = "Beta"
    ws["B3"] = "2"
    buf = io.BytesIO()
    wb.save(buf)
    wb.close()
    layout = parser.parse(buf.getvalue(), "data.xlsx")
    # No marker tokens → structural regions should be detected
    structural = [r for r in layout.table_regions if "structural" in r.region_id]
    assert len(structural) >= 1
    region = structural[0]
    assert region.header_row  # header row with Name / Value
