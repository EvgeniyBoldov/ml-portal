"""Unit tests for TemplateAnalyzeService."""
from __future__ import annotations

import pytest

from app.services.collection.template_analyze_service import TemplateAnalyzeService


@pytest.fixture
def service():
    return TemplateAnalyzeService()


@pytest.mark.asyncio
async def test_analyze_text_plain(service):
    content = b"Hello {{name}}, your code is {{code}}."
    result = await service.analyze_bytes(content, "greeting.txt")

    assert result["title"] == "Hello {{name}}, your code is {{code}}."
    placeholders = {item["placeholder"] for item in result["draft_schema"]["placeholders"]}
    assert placeholders == {"name", "code"}


@pytest.mark.asyncio
async def test_analyze_text_with_version(service):
    content = b"Template v2.1\nUser: {{user}}"
    result = await service.analyze_bytes(content, "report.txt")

    assert result["version"] == "2.1"
    assert {item["placeholder"] for item in result["draft_schema"]["placeholders"]} == {"user"}


@pytest.mark.asyncio
async def test_analyze_text_no_placeholders(service):
    content = b"Just a static text file."
    result = await service.analyze_bytes(content, "static.txt")

    assert result["title"] == "Just a static text file."
    assert result["draft_schema"]["placeholders"] == []
    assert result["version"] is None


@pytest.mark.asyncio
async def test_analyze_excel_mocked(service, monkeypatch):
    """Smoke-test that Excel path is reachable when openpyxl is present."""
    try:
        import openpyxl
    except ImportError:
        pytest.skip("openpyxl not installed")

    import io
    wb = openpyxl.Workbook()
    ws = wb.active
    ws["A1"] = "Name"
    ws["B1"] = "Value"
    ws["A2"] = "{{name}}"
    ws["B2"] = "{{amount}}"
    buf = io.BytesIO()
    wb.save(buf)
    wb.close()

    result = await service.analyze_bytes(buf.getvalue(), "data.xlsx")
    placeholders = {item["placeholder"] for item in result["draft_schema"]["placeholders"]}
    assert placeholders == {"name", "amount"}


@pytest.mark.asyncio
async def test_analyze_word_mocked(service, monkeypatch):
    """Smoke-test that Word path is reachable when python-docx is present."""
    try:
        import docx
    except ImportError:
        pytest.skip("python-docx not installed")

    import io
    doc = docx.Document()
    doc.add_paragraph("Hello {{customer}}")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "{{item}}"
    table.rows[0].cells[1].text = "{{price}}"
    buf = io.BytesIO()
    doc.save(buf)

    result = await service.analyze_bytes(buf.getvalue(), "letter.docx")
    placeholders = {item["placeholder"] for item in result["draft_schema"]["placeholders"]}
    assert placeholders == {"customer"}
    assert result["draft_schema"]["tables"] == [{"index": 0, "rows": 1, "columns": 2}]
