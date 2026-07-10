"""Tests for TemplateFillEngine (S4)."""
from __future__ import annotations
import io
import pytest
from app.services.collection.template_contract import (
    ScalarField, TableField, TableColumn, TemplateContract, FieldType,
    AnchorStrategy, TableAnchor, MarkerAnchor, StructuralAnchor,
)
from app.services.collection.template_fill_engine import TemplateFillEngine, FillResult


@pytest.fixture
def scalar_contract():
    return TemplateContract(fields=[
        ScalarField(key="name", label="Name", type=FieldType.STRING, required=True),
        ScalarField(key="amount", label="Amount", type=FieldType.NUMBER, required=True),
    ])


@pytest.fixture
def table_contract():
    return TemplateContract(fields=[
        ScalarField(key="company", label="Company", type=FieldType.STRING, required=True),
        TableField(
            key="items",
            label="Items",
            required=True,
            anchor=TableAnchor(
                strategy=AnchorStrategy.MARKER,
                marker=MarkerAnchor(loop_tokens=["{{items.name}}", "{{items.qty}}"]),
            ),
            columns=[
                TableColumn(key="name", label="Name", type=FieldType.STRING, required=True),
                TableColumn(key="qty", label="Qty", type=FieldType.NUMBER, required=True),
            ],
        ),
    ])


def test_fill_text_scalar_only(scalar_contract):
    engine = TemplateFillEngine(scalar_contract)
    template = b"Hello {{name}}, your amount is {{amount}}"
    values = {"name": "Alice", "amount": "100.50"}
    result = engine.fill(template, values, "test.txt")
    assert result.success is True
    assert b"Hello Alice" in result.content
    assert b"100.50" in result.content
    assert "name" in result.filled_scalars
    assert "amount" in result.filled_scalars


def test_fill_text_nested_and_typed_scalar():
    contract = TemplateContract(fields=[
        ScalarField(key="author.tel", label="Author tel", type=FieldType.NUMBER, required=True),
    ])
    engine = TemplateFillEngine(contract)
    template = b"Phone: {{author.tel(type=int,min=10)}}"
    result = engine.fill(template, {"author": {"tel": 12345}}, "test.txt")
    assert result.success is True
    assert b"12345" in result.content
    assert "author.tel" in result.filled_scalars


def test_fill_text_missing_required(scalar_contract):
    engine = TemplateFillEngine(scalar_contract)
    template = b"Hello {{name}}, amount {{amount}}"
    values = {"name": "Alice"}  # missing amount
    result = engine.fill(template, values, "test.txt")
    assert result.success is False
    assert "amount" in result.error.lower()
    assert "field required" in result.error.lower()


def test_fill_text_table_marker_loop(table_contract):
    engine = TemplateFillEngine(table_contract)
    template = b"{{#items}}{{items.name}} {{items.qty}}{{/items}}"
    values = {
        "company": "Acme",
        "items": [
            {"name": "Apple", "qty": 5},
            {"name": "Banana", "qty": 3},
        ]
    }
    result = engine.fill(template, values, "test.txt")
    assert result.success is True
    assert b"Apple 5" in result.content
    assert b"Banana 3" in result.content
    assert "items" in result.filled_tables


def test_fill_text_empty_required_table_fails(table_contract):
    engine = TemplateFillEngine(table_contract)
    template = b"{{#items}}{{items.name}}{{/items}}"
    values = {"company": "Acme", "items": []}
    result = engine.fill(template, values, "test.txt")
    assert result.success is False
    assert "required table 'items'" in result.error.lower()


def test_fill_text_fails_when_no_placeholders_were_replaced(scalar_contract):
    engine = TemplateFillEngine(scalar_contract)
    template = b"Static content without placeholders"
    values = {"name": "Alice", "amount": 100}
    result = engine.fill(template, values, "test.txt")
    assert result.success is False
    assert "no placeholders were filled" in result.error.lower()


def test_fill_text_empty_optional_table_succeeds():
    # An optional table with an empty list should pass validation and
    # produce empty content for the loop region.
    contract = TemplateContract(fields=[
        ScalarField(key="company", label="Company", type=FieldType.STRING, required=True),
        TableField(
            key="items", label="Items", required=False, min_rows=0,
            anchor=TableAnchor(
                strategy=AnchorStrategy.MARKER,
                marker=MarkerAnchor(loop_tokens=["{{items.name}}"]),
            ),
            columns=[TableColumn(key="name", label="Name", type=FieldType.STRING, required=True)],
        ),
    ])
    engine = TemplateFillEngine(contract)
    template = b"{{#items}}{{items.name}}{{/items}}"
    values = {"company": "Acme", "items": []}
    result = engine.fill(template, values, "test.txt")
    assert result.success is True
    assert result.content == b""


def test_validation_unknown_key_is_strict_error(scalar_contract):
    engine = TemplateFillEngine(scalar_contract)
    template = b"{{name}} {{amount}}"
    values = {"name": "Alice", "amount": 100, "unknown": "value"}
    result = engine.fill(template, values, "test.txt")
    assert result.success is False
    assert result.validation is not None
    assert any(issue.path == "unknown" for issue in result.validation.error_details)


def test_validation_type_mismatch(scalar_contract):
    engine = TemplateFillEngine(scalar_contract)
    template = b"test"
    values = {"name": "Alice", "amount": "not_a_number"}
    result = engine.fill(template, values, "test.txt")
    assert result.success is False
    assert "amount" in result.error.lower() or "number" in result.error.lower()
    assert result.validation is not None
    assert any(issue.path == "amount" for issue in result.validation.error_details)


def test_fill_excel_marker_loop_reuses_template_row():
    openpyxl = pytest.importorskip("openpyxl")

    contract = TemplateContract(fields=[
        TableField(
            key="items",
            label="Items",
            required=True,
            anchor=TableAnchor(
                strategy=AnchorStrategy.MARKER,
                marker=MarkerAnchor(loop_tokens=["{{items.name}}", "{{items.qty}}"]),
            ),
            columns=[
                TableColumn(key="name", label="Name", type=FieldType.STRING, required=True),
                TableColumn(key="qty", label="Qty", type=FieldType.NUMBER, required=True),
            ],
        ),
    ])
    wb = openpyxl.Workbook()
    ws = wb.active
    ws["A1"] = "Header"
    ws["A2"] = "{{items.name}}"
    ws["B2"] = "{{items.qty}}"
    ws.row_dimensions[2].height = 33
    stream = io.BytesIO()
    wb.save(stream)

    engine = TemplateFillEngine(contract)
    result = engine.fill(
        stream.getvalue(),
        {"items": [{"name": "Apple", "qty": 5}, {"name": "Banana", "qty": 3}]},
        "items.xlsx",
    )

    assert result.success is True
    loaded = openpyxl.load_workbook(io.BytesIO(result.content))
    out = loaded.active
    assert out["A2"].value == "Apple"
    assert out["B2"].value == "5"
    assert out["A3"].value == "Banana"
    assert out["B3"].value == "3"
    assert out.row_dimensions[2].height == 33
    assert out.row_dimensions[3].height == 33


def test_fill_excel_horizontal_marker_table_is_rejected():
    openpyxl = pytest.importorskip("openpyxl")

    contract = TemplateContract(fields=[
        TableField(
            key="items",
            label="Items",
            required=True,
            orientation="horizontal",
            anchor=TableAnchor(
                strategy=AnchorStrategy.MARKER,
                marker=MarkerAnchor(loop_tokens=["{{items.name}}"]),
            ),
            columns=[TableColumn(key="name", label="Name", type=FieldType.STRING, required=True)],
        ),
    ])
    wb = openpyxl.Workbook()
    ws = wb.active
    ws["A1"] = "{{items.name}}"
    stream = io.BytesIO()
    wb.save(stream)

    engine = TemplateFillEngine(contract)
    result = engine.fill(stream.getvalue(), {"items": [{"name": "Apple"}]}, "items.xlsx")

    assert result.success is False
    assert "horizontal table fill is not implemented yet" in result.error.lower()


def test_fill_docx_marker_loop_reuses_template_row():
    pytest.importorskip("docx")
    from docx import Document

    contract = TemplateContract(fields=[
        TableField(
            key="items",
            label="Items",
            required=True,
            anchor=TableAnchor(
                strategy=AnchorStrategy.MARKER,
                marker=MarkerAnchor(loop_tokens=["{{items.name}}", "{{items.qty}}"]),
            ),
            columns=[
                TableColumn(key="name", label="Name", type=FieldType.STRING, required=True),
                TableColumn(key="qty", label="Qty", type=FieldType.NUMBER, required=True),
            ],
        ),
    ])

    doc = Document()
    table = doc.add_table(rows=3, cols=2)
    table.rows[0].cells[0].text = "{{#items}}"
    table.rows[1].cells[0].text = "{{items.name}}"
    table.rows[1].cells[1].text = "{{items.qty}}"
    table.rows[2].cells[0].text = "{{/items}}"
    stream = io.BytesIO()
    doc.save(stream)

    engine = TemplateFillEngine(contract)
    result = engine.fill(
        stream.getvalue(),
        {"items": [{"name": "Apple", "qty": 5}, {"name": "Banana", "qty": 3}]},
        "items.docx",
    )

    assert result.success is True
    loaded = Document(io.BytesIO(result.content))
    out_table = loaded.tables[0]
    assert len(out_table.rows) == 2
    assert out_table.rows[0].cells[0].text == "Apple"
    assert out_table.rows[0].cells[1].text == "5"
    assert out_table.rows[1].cells[0].text == "Banana"
    assert out_table.rows[1].cells[1].text == "3"
